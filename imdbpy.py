import imdb
import requests
from bs4 import BeautifulSoup as bs

from docx import Document
from docx.shared import *



document=Document()

# Create the object that will be used to access the IMDb's database.
ia = imdb.IMDb() # by default access the web.

# Search for a movie (get a list of Movie objects).
m_result = ia.search_movie(input('enter a movie name:\n'))




# for item in m_result:
movie_id=m_result[0].movieID
url='http://www.imdb.com/title/tt'+str(movie_id)+'/'
  #print(url)
try:
 source=requests.get(url).text
 soup=bs(source,'lxml')
 html=soup.prettify()
 heading=soup.title.text
 document.add_heading(heading,0).bold = True
 bodi=soup.body.find('div',class_='summary_text')
 document.add_heading('Description:', level=1)

 document.add_paragraph(bodi.text.strip())

 #print('description:',bodi.text.strip())
 diractor=soup.body.find('span',class_='itemprop',itemprop='name')

#print('Director:',diractor.text)
 document.add_heading('Director:', level=1)
 document.add_paragraph(bodi.text.strip())

 rating=soup.body.find('span',itemprop='ratingValue')
#print('Rating:',rating.text+'/10')

 actors=soup.body.find('div',class_='article',id='titleCast').find_all('span',class_='itemprop',itemprop='name')
 charactors=soup.body.find('div',class_='article',id='titleCast').find_all('td',class_='character')

 print()

#print('Cast:')
 document.add_heading('Cast:', level=1)
 body=''
 for i in range(len(charactors)):
  body=body+actors[i].text.strip()+' as '+charactors[i].text.strip()+'\n'
 document.add_paragraph(body)
 print()
 story_head=soup.body.find('div',class_='article',id='titleStoryLine').find('h2')

#print(story_head.text,':\n')
 document.add_heading(story_head.text, level=1)

 story_body=soup.body.find('div',class_='article',id='titleStoryLine').find('div',class_='inline canwrap')

#print(story_body.text.strip())
 document.add_paragraph(story_body.text.strip())

 genres=soup.body.find('div',class_='article',id='titleStoryLine').find('div',class_='see-more inline canwrap',itemprop='genre').find_all('a')

#print('Genere:')
 document.add_heading('Genere:', level=1)
 for g in genres:
 #print(g.text,)
  document.add_paragraph(g.text)

 print()

 release_date=soup.body.find('div',class_='article',id='titleDetails').find_all('div',class_='txt-block')
 for i in range(2,len(release_date)-2):
  #print(release_date[i].text.replace('See more »','').replace('Show more on','').replace('  IMDbPro »','').strip())
  document.add_paragraph(release_date[i].text.replace('See more »','').replace('Show more on','').replace('  IMDbPro »','').strip())
 document.add_page_break()

 document.save('movie_data.docx')


except Exception as e:
    print (e)




